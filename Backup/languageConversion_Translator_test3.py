

from python_translator import Translator
import docx # python-docx==0.8.10

doc = docx.Document()

translator = Translator()

# For new document (document-wide):
# Set language value in the documents' default Run's Properties element.
styles_element = doc.styles.element
print("----------------------------------------------------------")
print(styles_element)
rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
print("----------------------------------------------------------")
print(rpr_default)
lang_default = rpr_default.xpath('w:lang')[0]
print("----------------------------------------------------------")
print(lang_default)
lang_default.set(docx.oxml.shared.qn('w:val'),'jp')



p4 = doc.add_paragraph(
    '三星ナナミ',
    style='Normal'
)

result = translator.translate("三星ナナミ", "english", "japanese")
#result = translator.translate("fuck", "japanese", "english")

print("----------------------------------------------------------")
print(result)
print(type(result))
print(dir(result))


print("----------------------------------------------------------")
print("new translated text.")
print(result.new_text)
print(type(result.new_text))
print(dir(result.new_text))



p4 = doc.add_paragraph(
    result.new_text,
    # style='Normal'
)


doc.save('my-document.docx')