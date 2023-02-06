from __future__ import absolute_import, division, print_function, unicode_literals

import docx

from docx import Document
from docx.shared import Inches

from youtube_transcript_api import YouTubeTranscriptApi

import json

from googletrans import Translator

import scrapetube




def YT_Video_Transcript(video_url,document_name,document_ext,transcript_language):
    
    translator = Translator()
    document = Document()

    
    #section._Footer.is_linked_to_previous=True
    
    video_id=video_url.split('=')[1]
    print(video_id)
    video_info=YouTubeTranscriptApi.get_transcript(video_id,languages=transcript_language)
    print('-----------------------')
    
    line=int(input("How many line do you want to handle with?"))
    msg="The transcript of YT, link="+str(video_url)
    print(msg)
    document.add_paragraph(msg)
    document.add_page_break()
    count=0
    
    for v_info in video_info:
        count+=1
        for la in ['zh-tw','en','ja']:
            backup=v_info
            msg_para="The current language is "+str(la)
            print(msg_para)
            document.add_paragraph(msg_para)
            
            text=v_info['text']
            x=translator.translate(text, dest=la)
            
            backup['text']=x.text
            s=str(backup)
            print(s)
            document.add_paragraph(s)
        
        if count>=line:
            break
    document.save(document_name+str(document_ext))


def main():
    
    videos = scrapetube.get_channel("UCO-n4ZDDXKPKK29c5eaytpA")
    base_watch_video_url='https://www.youtube.com/watch?v='

    cnt=0
    for video in videos:
        cnt+=1
        v_url=video['videoId']
        video_url=str(base_watch_video_url)+v_url
        print(video_url)
        YT_Video_Transcript(video_url,'Nanaten'+str(cnt),'.docx',['ja'])

if __name__=='__main__':
    main()