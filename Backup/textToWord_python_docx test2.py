from docx import Document
from docx.shared import Inches
from youtube_transcript_api import YouTubeTranscriptApi
import json
from googletrans import Translator


translator = Translator()

video_url='https://www.youtube.com/watch?v=JMOp5n3YgLA'
video_id=video_url.split('=')[1]
print(video_id)
video_info=YouTubeTranscriptApi.get_transcript(video_id)
print('-----------------------')
    
document = Document()


line=int(input("How many line do you want to handle with?"))
msg="The transcript of YT, link="+str(video_url)
print(msg)
document.add_paragraph(msg)
document.add_page_break()
count=0

for v_info in video_info:
    count+=1
    for la in ['zh-tw','en','ja']:
        backup=v_info
        msg_para="The current language is "+str(la)
        print(msg_para)
        document.add_paragraph(msg_para)
        
        text=v_info['text']
        x=translator.translate(text, dest=la)
    
        backup['text']=x.text
        
        s=str(backup)
        
        #print(type(s))
        print(s)
        
        document.add_paragraph(s)
        
    if count>=line:
        break
document.save('demo_multilanguage.docx')