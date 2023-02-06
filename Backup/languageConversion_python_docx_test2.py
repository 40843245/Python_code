
import docx # python-docx==0.8.10
from python_translator import Translator

translator = Translator()
doc = docx.Document()

# For new document (document-wide):
# Set language value in the documents' default Run's Properties element.
styles_element = doc.styles.element
print("----------------------------------------------------------")
print(styles_element)
rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
print("----------------------------------------------------------")
print(rpr_default)
lang_default = rpr_default.xpath('w:lang')[0]
print("----------------------------------------------------------")
print(lang_default)
lang_default.set(docx.oxml.shared.qn('w:val'),'jp')

p4 = doc.add_paragraph(
    '三星ナナミ',
    style='Normal'
)

result=translator.translate("fuck", "japanese", "english")
p4 = doc.add_paragraph(
    result.new_text,
    # style='Normal'
)



doc.save('my-document.docx')